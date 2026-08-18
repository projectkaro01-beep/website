import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_report():
    doc = docx.Document()

    # Page Margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base typography styling
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x2D, 0x37, 0x48) # Slate dark

    # Helper function for cell background color
    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Helper function for cell margins/padding
    def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # =========================================================================
    # TITLE & HEADER SECTION
    # =========================================================================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_title.add_run("COLLEGE CYBERSECURITY & WEB APPLICATION SECURITY HACKATHON 2026\n")
    r_sub.font.size = Pt(10)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0) # Blue

    r_title = p_title.add_run("ADVANCED PENETRATION TESTING & REMEDIATION REPORT\n")
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

    r_proj = p_title.add_run("Full-Stack Student Management System — Complete Security Lifecycle")
    r_proj.font.size = Pt(12.5)
    r_proj.font.italic = True
    r_proj.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    meta_data = [
        ("Team Name:", "Bug Smasher"),
        ("Team Leader:", "Ranjan Gowda S S"),
        ("Team Members:", "1. Ranjan Gowda S S (Team Leader — Security & Full-Stack Architect)\n2. Manish M (Lead Penetration Tester & Exploit Researcher)\n3. Keerthan Gowda T R (Database Administrator & Hardening Engineer)"),
        ("Audit & Remediation Scope:", "10 Vulnerabilities Discovered → Controlled Exploits Verified → 100% Remediated via Defense-in-Depth in Secure V2")
    ]

    col_widths = [Inches(1.8), Inches(4.7)]
    for r_idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[r_idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.width = col_widths[0]
        cell_val.width = col_widths[1]
        
        set_cell_background(cell_lbl, "F7FAFC")
        set_cell_background(cell_val, "FFFFFF")
        set_cell_margins(cell_lbl, top=80, bottom=80, left=120, right=120)
        set_cell_margins(cell_val, top=80, bottom=80, left=120, right=120)

        p_l = cell_lbl.paragraphs[0]
        r_l = p_l.add_run(label)
        r_l.bold = True
        r_l.font.size = Pt(9.5)
        r_l.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

        p_v = cell_val.paragraphs[0]
        r_v = p_v.add_run(val)
        r_v.font.size = Pt(9.5)
        if "100% Remediated" in val:
            r_v.bold = True
            r_v.font.color.rgb = RGBColor(0x27, 0x67, 0x49)

    doc.add_paragraph()

    # =========================================================================
    # SECTION 1: EXECUTIVE SUMMARY
    # =========================================================================
    h1 = doc.add_heading("1. Executive Summary & Assessment Methodology", level=1)
    h1.style.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    p_intro = doc.add_paragraph()
    p_intro.add_run(
        "Modern web applications built on Single Page Application (SPA) frameworks coupled directly with Backend-as-a-Service (BaaS) platforms like Supabase introduce significant security risks when developers confuse user interface routing with server-side authorization. When database tables lack fine-grained Row Level Security (RLS) or when input parameters are unvalidated, low-privilege users can bypass the visual frontend entirely and execute unauthorized operations directly against the database API.\n\n"
        "Team "
    )
    p_intro.add_run("Bug Smasher").bold = True
    p_intro.add_run(
        " executed a comprehensive, dual-phase security lifecycle:\n"
        "• Phase 1 (Target Construction & Attack Verification): Constructed a functional, multi-tier Student Management System (EduPulse) containing authentic PostgreSQL relationships, Supabase Authentication, real role models (Student, Teacher, Admin), and genuine academic workflows. We identified, reproduced, and documented 10 controlled vulnerabilities spanning access control, injection, parameter tampering, and storage isolation.\n"
        "• Phase 2 (Defense-in-Depth Engineering): Engineered an enterprise-grade remediation suite that moves beyond cosmetic UI fixes. We implemented PostgreSQL SECURITY DEFINER role validation functions, database-level BEFORE UPDATE mutation triggers, strict row-ownership RLS predicates, client-side input whitelisting, DOMPurify HTML sanitization, and ephemeral 60-second storage signed tokens."
    )

    # =========================================================================
    # SECTION 2: ATTACK MATRIX & SUMMARY TABLE
    # =========================================================================
    h2 = doc.add_heading("2. Threat Modeling & Vulnerability Summary Matrix", level=1)
    h2.style.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    summary_table = doc.add_table(rows=11, cols=6)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.autofit = False

    headers = ["ID", "Vulnerability Title", "OWASP Category", "Severity", "Exploited Attack Result", "Secure V2 Prevention"]
    hdr_widths = [Inches(0.8), Inches(1.8), Inches(1.3), Inches(0.8), Inches(1.1), Inches(0.8)]

    for c_idx, head_text in enumerate(headers):
        cell = summary_table.rows[0].cells[c_idx]
        cell.width = hdr_widths[c_idx]
        set_cell_background(cell, "1A365D")
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        p = cell.paragraphs[0]
        r = p.add_run(head_text)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    matrix_rows = [
        ("VULN-001", "Client-Side Route Guarding", "Broken Access Control", "HIGH", "Injected courses as Student", "BLOCKED"),
        ("VULN-002", "Insecure Direct Object Ref (IDOR)", "Broken Access Control", "HIGH", "Accessed Student B's grades", "BLOCKED"),
        ("VULN-003", "RLS Misconfiguration", "Security Misconfig", "HIGH", "Dumped all marks/attendance", "BLOCKED"),
        ("VULN-004", "Privilege Escalation via Role", "Broken Access Control", "CRITICAL", "Student elevated to Admin", "BLOCKED"),
        ("VULN-005", "Stored Cross-Site Scripting (XSS)", "Injection", "HIGH", "Executed JS in notice board", "BLOCKED"),
        ("VULN-006", "Mass Assignment Parameter Binding", "Data Integrity", "HIGH", "Mutated backend attributes", "BLOCKED"),
        ("VULN-007", "Storage Cross-Tenant Exposure", "Broken Access Control", "HIGH", "Downloaded private files", "BLOCKED"),
        ("VULN-008", "Sensitive Data Exposure (PII)", "Data Exposure", "MEDIUM", "Harvested student phone numbers", "BLOCKED"),
        ("VULN-009", "Attendance Logic Bypass", "Insecure Design", "HIGH", "Falsified attendance to Present", "BLOCKED"),
        ("VULN-010", "Debug State Information Leak", "Security Misconfig", "MEDIUM", "Extracted tokens via console", "BLOCKED"),
    ]

    for r_idx, r_data in enumerate(matrix_rows):
        row = summary_table.rows[r_idx + 1]
        bg_color = "F7FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            cell.width = hdr_widths[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=70, bottom=70, left=70, right=70)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8)
            
            if c_idx == 3:
                r.bold = True
                if val == "CRITICAL":
                    r.font.color.rgb = RGBColor(0x9B, 0x2C, 0x2C)
                elif val == "HIGH":
                    r.font.color.rgb = RGBColor(0xC5, 0x30, 0x30)
                else:
                    r.font.color.rgb = RGBColor(0xC0, 0x56, 0x21)
            elif c_idx == 5:
                r.bold = True
                r.font.color.rgb = RGBColor(0x27, 0x67, 0x49)

    doc.add_paragraph()

    # =========================================================================
    # SECTION 3: TECHNICAL AUDIT, ATTACKS & EXHAUSTIVE PREVENTIONS
    # =========================================================================
    h3 = doc.add_heading("3. Detailed Vulnerability Findings, Attack Reproductions & Prevention Architectures", level=1)
    h3.style.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    vuln_details = [
        {
            "id": "VULN-001",
            "title": "Broken Access Control — Client-Side Only Route & Action Guarding",
            "severity": "HIGH (CVSS: 8.1)",
            "category": "OWASP A01:2021 — Broken Access Control",
            "target": "Course Curriculum & Faculty Directory (public.subjects, public.teachers)",
            "attack_desc": "In modern React single-page applications, UI route guards (such as ProtectedRoute) only alter client-side browser views and do not represent a security barrier. In Vulnerable V1, while administrative buttons were hidden from the Student UI, the underlying Supabase tables lacked backend role enforcement. Under an authenticated Student session, we executed direct API queries via DevTools Console to insert new course records into public.subjects and extract faculty records.",
            "attack_res": "The API accepted the request with HTTP 200 OK, creating course 'HACK-101' and dumping the full instructor database directly from a standard student session.",
            "impact": "Complete collapse of role segregation. Any registered student could inject fictitious courses, tamper with curriculum structures, or corrupt administrative tables.",
            "screenshot_box": "VULN-001-before.png\nCaption: DevTools console showing student account successfully executing administrative course insertion into public.subjects.",
            "prevention_deep_dive": (
                "A. Database-Level Role Evaluation Functions:\n"
                "Relying on JWT user metadata for authorization is vulnerable to token tampering and synchronization delays. We implemented PostgreSQL helper functions with SECURITY DEFINER to inspect the verified public.profiles record:\n\n"
                "CREATE OR REPLACE FUNCTION public.is_admin()\n"
                "RETURNS boolean AS $$\n"
                "  SELECT EXISTS (SELECT 1 FROM public.profiles WHERE id = auth.uid() AND role = 'admin');\n"
                "$$ LANGUAGE sql SECURITY DEFINER STABLE;\n\n"
                "B. Hardened Row Level Security (RLS) Policy:\n"
                "We locked all modification operations on public.subjects and public.teachers exclusively to administrators:\n\n"
                "CREATE POLICY \"subjects_admin_manage_policy\" ON public.subjects\n"
                "FOR ALL TO authenticated\n"
                "USING (public.is_admin())\n"
                "WITH CHECK (public.is_admin());\n\n"
                "C. Verification & Failure State:\n"
                "Any direct API mutation issued by a non-admin session is intercepted by PostgreSQL before reaching the table buffer, returning an immediate fatal RLS permission violation (HTTP 403 Forbidden)."
            )
        },
        {
            "id": "VULN-002",
            "title": "Insecure Direct Object Reference (IDOR / BOLA) on Academic Marks",
            "severity": "HIGH (CVSS: 8.5)",
            "category": "OWASP A01:2021 — Broken Access Control / API1:2023 BOLA",
            "target": "Confidential Grade Ledger & Academic Evaluations of Student B",
            "attack_desc": "The marks lookup component accepted an untrusted query parameter (?student_id=<UUID>). Because the backend failed to validate whether the requesting user owned the requested student_id, an attacker could supply another student's identifier in the URL or lookup bar to retrieve their grade report.",
            "attack_res": "Student A successfully rendered Student B's complete grade sheet, including Midterm/Final examination scores, quiz grades, percentages, and private faculty remarks.",
            "impact": "Severe breach of academic confidentiality and data protection laws (FERPA, GDPR). Every student could spy on the marks and academic standing of all peers.",
            "screenshot_box": "VULN-002-before.png\nCaption: Student A's interface displaying Student B's confidential academic marks sheet and teacher remarks via IDOR parameter tampering.",
            "prevention_deep_dive": (
                "A. Elimination of Client-Controlled Query Identifiers:\n"
                "In StudentMarks.tsx, we completely removed URL parameter binding and manual search lookups. The client query is bound strictly to the session's internal studentRecord.id resolved on the server side.\n\n"
                "B. Identity-Scoped Row Level Security (RLS):\n"
                "We enforced strict ownership predicates in PostgreSQL ensuring students can ONLY select rows matching their own authenticated student profile:\n\n"
                "CREATE OR REPLACE FUNCTION public.get_student_id()\n"
                "RETURNS UUID AS $$\n"
                "  SELECT id FROM public.students WHERE user_id = auth.uid();\n"
                "$$ LANGUAGE sql SECURITY DEFINER STABLE;\n\n"
                "CREATE POLICY \"marks_select_policy\" ON public.marks FOR SELECT TO authenticated\n"
                "USING (\n"
                "    student_id = public.get_student_id()\n"
                "    OR public.is_admin()\n"
                "    OR (public.is_teacher() AND subject_id IN (\n"
                "        SELECT id FROM public.subjects WHERE teacher_id = public.get_teacher_id()\n"
                "    ))\n"
                ");\n\n"
                "C. Verification & Failure State:\n"
                "If Student A attempts to query Student B's marks via manual API request, PostgreSQL filters the query using the WHERE predicate student_id = public.get_student_id(), returning an empty array (0 rows) and preventing unauthorized data discovery."
            )
        },
        {
            "id": "VULN-003",
            "title": "PostgreSQL Row Level Security (RLS) Misconfiguration — Overly Permissive Policies",
            "severity": "HIGH (CVSS: 8.2)",
            "category": "OWASP A05:2021 — Security Misconfiguration",
            "target": "Entire Institutional Marks & Attendance Database",
            "attack_desc": "Database tables had RLS enabled, but policies were configured with open clauses: FOR SELECT TO authenticated USING (true). This granted any authenticated bearer token blanket read access across every record in the table regardless of role or ownership.",
            "attack_res": "Executing an unconstrained SELECT * against public.marks and public.attendance dumped the complete institution-wide academic ledger and attendance history across all students in a single query.",
            "impact": "Unrestricted bulk data exfiltration. Any registered user could extract thousands of institutional records directly via REST endpoints.",
            "screenshot_box": "VULN-003-before.png\nCaption: Browser console table displaying an unfiltered dump of all students' marks and attendance logs via misconfigured RLS policies.",
            "prevention_deep_dive": (
                "A. Elimination of Open 'USING (true)' Directives:\n"
                "In PostgreSQL BaaS architectures, a policy with USING (true) completely bypasses tenant isolation. We dropped all open policies across profiles, students, teachers, marks, attendance, and documents.\n\n"
                "B. Role-Based & Tenant-Scoped Predicates:\n"
                "Every table was partitioned with mathematical tenant guarantees:\n"
                "• public.profiles: id = auth.uid() OR public.is_admin() OR public.is_teacher()\n"
                "• public.attendance: student_id = public.get_student_id() OR public.is_admin() OR assigned_teacher\n"
                "• public.marks: student_id = public.get_student_id() OR public.is_admin() OR assigned_teacher\n\n"
                "C. Verification & Failure State:\n"
                "When a student issues an unconstrained SELECT * FROM marks, the database runtime automatically injects the active user's RLS constraints into the query execution planner, ensuring only rows owned by the caller are ever evaluated."
            )
        },
        {
            "id": "VULN-004",
            "title": "Privilege Escalation — Unauthorized Role Field Mutation",
            "severity": "CRITICAL (CVSS: 9.8)",
            "category": "OWASP A01:2021 — Broken Access Control / Privilege Escalation",
            "target": "User Profiles & System Authorization Table (public.profiles.role)",
            "attack_desc": "The profiles table contained the role column (student, teacher, admin). The update RLS policy allowed users to update their own row without column-level restrictions. Under a Student account, we issued a direct update payload { role: 'admin' } and refreshed the application.",
            "attack_res": "The database committed the role update. Upon page refresh, the application re-hydrated the session with super-admin privileges, dynamically unlocking the full ADMIN PORTAL and administrative management modules.",
            "impact": "Complete system takeover. Any unprivileged user could unilaterally promote their account to super-administrator and assume total administrative control over the entire institution.",
            "screenshot_box": "VULN-004-before.png\nCaption: Student A's account interface immediately following the attack, displaying the red ADMIN role badge and Admin Portal sidebar navigation.",
            "prevention_deep_dive": (
                "A. Immutable Role Field Trigger (Database Layer):\n"
                "Because column-level UPDATE grants in PostgreSQL RLS can still be bypassed if policy checks are too broad, we implemented an un-bypassable BEFORE UPDATE trigger on public.profiles:\n\n"
                "CREATE OR REPLACE FUNCTION public.protect_profile_role()\n"
                "RETURNS TRIGGER AS $$\n"
                "BEGIN\n"
                "    IF NEW.role IS DISTINCT FROM OLD.role THEN\n"
                "        IF NOT public.is_admin() THEN\n"
                "            RAISE EXCEPTION 'Unauthorized: Only administrators can modify user roles.';\n"
                "        END IF;\n"
                "    END IF;\n"
                "    RETURN NEW;\n"
                "END;\n"
                "$$ LANGUAGE plpgsql SECURITY DEFINER;\n\n"
                "CREATE TRIGGER trg_protect_profile_role\n"
                "    BEFORE UPDATE ON public.profiles\n"
                "    FOR EACH ROW EXECUTE FUNCTION public.protect_profile_role();\n\n"
                "B. Server-Side Role Enforcement:\n"
                "Even if a malicious actor sends { role: 'admin' } in raw JSON, the database engine aborts the transaction before disk write, throwing an uncatchable PL/pgSQL exception."
            )
        },
        {
            "id": "VULN-005",
            "title": "Stored Cross-Site Scripting (XSS) on Campus Notice Board",
            "severity": "HIGH (CVSS: 8.3)",
            "category": "OWASP A03:2021 — Injection (Stored XSS)",
            "target": "Campus Bulletin Board & Victim Browser Sessions (Students/Teachers)",
            "attack_desc": "Announcements posted to the notice board were rendered into the DOM using React's unescaped dangerouslySetInnerHTML directive without HTML sanitization. We posted a notice containing an event-driven payload: <img src=invalid onerror=\"alert('VULN-005: Stored XSS Executed on ' + document.domain);\" />.",
            "attack_res": "When Student A viewed the notice board, the browser parsed the unescaped tag, triggered the onerror handler, and executed arbitrary JavaScript in the victim's session context.",
            "impact": "Session hijacking (exfiltrating JWT tokens from storage), credential harvesting via DOM phishing overlays, and forced background administrative actions.",
            "screenshot_box": "VULN-005-before.png\nCaption: Browser alert dialog popping up over the Student Notice Board confirming arbitrary script execution via unescaped stored HTML.",
            "prevention_deep_dive": (
                "A. DOMPurify Sanitization Engine Integration:\n"
                "We integrated DOMPurify across all notice rendering components (StudentNotices.tsx, StudentDashboard.tsx, TeacherNotices.tsx, AdminNotices.tsx). Every HTML fragment is passed through a strict sanitization pipeline prior to DOM injection:\n\n"
                "import DOMPurify from 'dompurify';\n\n"
                "const sanitizedContent = DOMPurify.sanitize(notice.content, {\n"
                "  USE_PROFILES: { html: true },\n"
                "  FORBID_TAGS: ['script', 'iframe', 'object', 'embed', 'form'],\n"
                "  FORBID_ATTR: ['onerror', 'onload', 'onclick', 'onmouseover', 'onfocus'],\n"
                "});\n\n"
                "B. Neutralization of Executable Vectors:\n"
                "When malicious payloads such as <img src=x onerror=alert(1)> are processed, DOMPurify strips the onerror handler and renders a safe, harmless <img> tag, completely disarming execution while maintaining valid formatting."
            )
        },
        {
            "id": "VULN-006",
            "title": "Mass Assignment / Unvalidated Parameter Binding",
            "severity": "HIGH (CVSS: 7.5)",
            "category": "OWASP A08:2021 — Software & Data Integrity Failures",
            "target": "Backend User Model & Protected Database Properties",
            "attack_desc": "The profile update routine accepted arbitrary client objects and forwarded them directly into the database update query without filtering or schema whitelisting. We submitted a payload containing extra, unexposed fields (role: 'teacher', avatar_url, updated_at).",
            "attack_res": "The database bound and committed all supplied keys into the profile record in a single transaction without parameter rejection.",
            "impact": "Attackers can manipulate hidden model attributes, overwrite audit timestamps, or modify privilege flags during standard profile saves.",
            "screenshot_box": "VULN-006-before.png\nCaption: Console showing unvalidated mass assignment payload updating multiple internal profile attributes simultaneously.",
            "prevention_deep_dive": (
                "A. Strict DTO Field Whitelisting (Frontend & API Layer):\n"
                "In StudentProfile.tsx, we replaced direct object ingestion with an explicit Data Transfer Object (DTO) constructor that extracts only explicitly approved user-modifiable properties:\n\n"
                "const sanitizedPayload = {\n"
                "  full_name: fullName.trim(),\n"
                "  phone: phone.trim(),\n"
                "  updated_at: new Date().toISOString(),\n"
                "};\n"
                "await supabase.from('profiles').update(sanitizedPayload).eq('id', profile.id);\n\n"
                "B. Database Layer Constraints:\n"
                "Combined with the PostgreSQL role protection trigger, any client-side attempt to attach unapproved fields (such as role, created_at, or status) is discarded by the client constructor and rejected by the database."
            )
        },
        {
            "id": "VULN-007",
            "title": "Insecure Storage Bucket Access Control & Cross-Tenant File Exposure",
            "severity": "HIGH (CVSS: 7.9)",
            "category": "OWASP A01:2021 — Broken Access Control",
            "target": "Private Student Verification Files & Storage Vault (student-documents)",
            "attack_desc": "The storage bucket was configured as public = true, and read policies allowed unconstrained reads. Student A uploaded a private document. Student B queried the document registry, retrieved Student A's storage key (<UUID>/1700000000-assignment.pdf), and generated the public storage URL.",
            "attack_res": "The storage endpoint returned Student A's document with HTTP 200 OK directly to Student B and unauthenticated users in an Incognito window.",
            "impact": "Confidentiality loss of student identity documents, medical records, and assignment submissions across tenants.",
            "screenshot_box": "VULN-007-before.png\nCaption: Storage Explorer showing Student B retrieving Student A's storage key, with an adjacent incognito tab downloading the file unauthenticated.",
            "prevention_deep_dive": (
                "A. Private Bucket Transition & Storage RLS:\n"
                "We set the storage bucket to private (public = false) and implemented strict RLS on storage.objects requiring upload/read folder paths to match the user's authentic auth.uid():\n\n"
                "CREATE POLICY \"storage_user_upload_policy\" ON storage.objects FOR INSERT TO authenticated\n"
                "WITH CHECK (\n"
                "    bucket_id = 'student-documents' \n"
                "    AND (storage.foldername(name))[1] = auth.uid()::text\n"
                ");\n\n"
                "CREATE POLICY \"storage_user_read_policy\" ON storage.objects FOR SELECT TO authenticated\n"
                "USING (\n"
                "    bucket_id = 'student-documents' \n"
                "    AND ((storage.foldername(name))[1] = auth.uid()::text OR public.is_admin() OR public.is_teacher())\n"
                ");\n\n"
                "B. Ephemeral Signed URL Token Generation:\n"
                "In StudentDocuments.tsx, files are never downloaded via static public links. The application issues a secure call generating a cryptographic, 60-second signed URL:\n\n"
                "const { data, error } = await supabase.storage\n"
                "  .from('student-documents')\n"
                "  .createSignedUrl(filePath, 60);\n\n"
                "Unauthenticated or cross-tenant access attempts receive an immediate HTTP 403 Access Denied error."
            )
        },
        {
            "id": "VULN-008",
            "title": "Sensitive Data Exposure in Roster & Directory Queries",
            "severity": "MEDIUM (CVSS: 5.3)",
            "category": "OWASP A02:2021 — Sensitive Data Exposure / Broken Property Authorization",
            "target": "Personal Identifiable Information (PII) of Enrolled Students",
            "attack_desc": "Roster queries used wildcard selections (students(*, profile:profiles(*))), over-fetching all columns from the database and delivering private contact phone numbers and internal UUIDs into client memory.",
            "attack_res": "Auditing the network response payload revealed complete profile structures containing private phone numbers and internal user IDs for all enrolled students.",
            "impact": "Violation of Data Minimization (GDPR Article 5). Exposes students' private contact numbers, increasing the risk of targeted social engineering and phishing.",
            "screenshot_box": "VULN-008-before.png\nCaption: Network tab showing full profile structures including private phone numbers returned in student roster response payload.",
            "prevention_deep_dive": (
                "A. Query Projection Data Minimization:\n"
                "In TeacherStudents.tsx and StudentSubjects.tsx, we replaced wildcard selections with explicit, minimal column projections:\n\n"
                "const { data } = await supabase\n"
                "  .from('enrollments')\n"
                "  .select(`\n"
                "    student:students(\n"
                "      id,\n"
                "      enrollment_no,\n"
                "      department,\n"
                "      semester,\n"
                "      profile:profiles(full_name, email)\n"
                "    ),\n"
                "    subject:subjects(name, code)\n"
                "  `);\n\n"
                "B. Privacy & Principle of Least Privilege:\n"
                "Private phone numbers and raw foreign keys are excluded from the API response payload, ensuring sensitive student PII is never transmitted to unauthorized client viewports."
            )
        },
        {
            "id": "VULN-009",
            "title": "Business Logic Authorization Failure — Attendance Self-Modification",
            "severity": "HIGH (CVSS: 7.7)",
            "category": "OWASP A04:2021 — Insecure Design / Broken Access Control",
            "target": "Institutional Classroom Attendance Records (public.attendance)",
            "attack_desc": "The application assumed hiding the attendance marking UI was sufficient to restrict attendance marking to instructors. The database policy allowed all authenticated users to execute INSERT and UPDATE on public.attendance. Under a Student session with status Absent, we issued a direct upsert call marking Student A as Present.",
            "attack_res": "The database committed the modification. Upon refreshing the student portal, Student A's attendance reflected Present and their attendance percentage increased immediately.",
            "impact": "Total loss of institutional academic integrity. Students can falsify classroom attendance records, bypass exam eligibility requirements, and overwrite faculty records.",
            "screenshot_box": "VULN-009-before.png\nCaption: Student Attendance view showing attendance percentage and session status successfully falsified to Present via student session API injection.",
            "prevention_deep_dive": (
                "A. Business Logic Verification via Database RLS:\n"
                "We engineered the PostgreSQL RLS policy attendance_modify_policy to strictly validate that the caller is either an administrator or the specific faculty member assigned to instruct that subject:\n\n"
                "CREATE POLICY \"attendance_modify_policy\" ON public.attendance\n"
                "FOR ALL TO authenticated\n"
                "USING (\n"
                "    public.is_admin() \n"
                "    OR (\n"
                "        public.is_teacher() AND subject_id IN (\n"
                "            SELECT id FROM public.subjects WHERE teacher_id = public.get_teacher_id()\n"
                "        )\n"
                "    )\n"
                ")\n"
                "WITH CHECK (\n"
                "    public.is_admin() \n"
                "    OR (\n"
                "        public.is_teacher() AND subject_id IN (\n"
                "            SELECT id FROM public.subjects WHERE teacher_id = public.get_teacher_id()\n"
                "        )\n"
                "    )\n"
                ");\n\n"
                "B. Verification & Failure State:\n"
                "If a student attempts to execute an attendance upsert, the database policy check evaluates to false and rejects the mutation with an RLS access violation error."
            )
        },
        {
            "id": "VULN-010",
            "title": "Security Misconfiguration — Global Debug Object & Client State Exposure",
            "severity": "MEDIUM (CVSS: 4.8)",
            "category": "OWASP A05:2021 — Security Misconfiguration",
            "target": "Client Runtime Environment & Active Authentication Session",
            "attack_desc": "An internal diagnostic object was bound directly to window.__APP_DEBUG__, providing helper methods that returned active session JWT tokens, project URLs, and client handles from the browser console.",
            "attack_res": "Typing window.__APP_DEBUG__.getTokens() into DevTools instantly extracted the user's active session tokens and configuration without inspecting local storage or cookies.",
            "impact": "Facilitates rapid reconnaissance and streamlines script-assisted session token harvesting.",
            "screenshot_box": "VULN-010-before.png\nCaption: Browser console displaying the exposed window.__APP_DEBUG__ object and token extraction helper output.",
            "prevention_deep_dive": (
                "A. Complete Removal of Diagnostic Interfaces:\n"
                "In src/lib/supabase.ts, all global window object attachments, debug harnesses, and token extraction functions were completely removed:\n\n"
                "// Secure V2: Diagnostic harnesses stripped from production bundle\n"
                "export const supabase = createClient(supabaseUrl, supabaseAnonKey);\n\n"
                "B. Verification & Failure State:\n"
                "Attempting to query window.__APP_DEBUG__ in DevTools returns undefined, preventing automated client-side reconnaissance and token scraping."
            )
        }
    ]

    for v in vuln_details:
        h_v = doc.add_heading(f"{v['id']}: {v['title']}", level=2)
        h_v.style.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

        # Key details bullet list
        p_k = doc.add_paragraph()
        r_sev_lbl = p_k.add_run("• Severity Rating: ")
        r_sev_lbl.bold = True
        r_sev = p_k.add_run(v['severity'])
        r_sev.bold = True
        if "CRITICAL" in v['severity']:
            r_sev.font.color.rgb = RGBColor(0x9B, 0x2C, 0x2C)
        elif "HIGH" in v['severity']:
            r_sev.font.color.rgb = RGBColor(0xC5, 0x30, 0x30)
        else:
            r_sev.font.color.rgb = RGBColor(0xC0, 0x56, 0x21)

        p_k.add_run(f"\n• OWASP Category: {v['category']}")
        p_k.add_run(f"\n• Target Asset: {v['target']}")

        # Attack Breakdown
        p_a = doc.add_paragraph()
        r_ab = p_a.add_run("1. Technical Root Cause & Attack Methodology:\n")
        r_ab.bold = True
        p_a.add_run(v['attack_desc'])

        p_res = doc.add_paragraph()
        r_rb = p_res.add_run("2. Attack Outcome & Verification:\n")
        r_rb.bold = True
        p_res.add_run(v['attack_res'])

        p_imp = doc.add_paragraph()
        r_ib = p_imp.add_run("3. Business Impact & Risk Analysis:\n")
        r_ib.bold = True
        p_imp.add_run(v['impact'])

        # SCREENSHOT PLACEHOLDER BOX (Bordered Box)
        p_box_lbl = doc.add_paragraph()
        r_box_lbl = p_box_lbl.add_run("4. Proof-of-Concept Screenshot Evidence:")
        r_box_lbl.bold = True

        box_table = doc.add_table(rows=1, cols=1)
        box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        box_cell = box_table.rows[0].cells[0]
        box_cell.width = Inches(6.5)
        set_cell_background(box_cell, "F7FAFC")
        set_cell_margins(box_cell, top=180, bottom=180, left=180, right=180)
        
        tcPr = box_cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="12" w:space="0" w:color="CBD5E0"/><w:bottom w:val="single" w:sz="12" w:space="0" w:color="CBD5E0"/><w:left w:val="single" w:sz="12" w:space="0" w:color="CBD5E0"/><w:right w:val="single" w:sz="12" w:space="0" w:color="CBD5E0"/></w:tcBorders>')
        tcPr.append(tcBorders)

        p_inner = box_cell.paragraphs[0]
        p_inner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_in1 = p_inner.add_run(f"[ PASTE SCREENSHOT HERE: {v['id']}-before.png ]\n\n")
        r_in1.bold = True
        r_in1.font.size = Pt(11)
        r_in1.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
        
        r_in2 = p_inner.add_run(v['screenshot_box'].split('\n')[1])
        r_in2.font.size = Pt(9.5)
        r_in2.font.italic = True
        r_in2.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

        # Exhaustive Prevention Deep-Dive
        p_def = doc.add_paragraph()
        r_db = p_def.add_run("5. Exhaustive Prevention Architecture (Secure V2 Remediation):\n")
        r_db.bold = True
        r_db.font.color.rgb = RGBColor(0x27, 0x67, 0x49)
        p_def.add_run(v['prevention_deep_dive'])

        doc.add_paragraph()

    # =========================================================================
    # SECTION 4: RETESTING & VERIFICATION MATRIX
    # =========================================================================
    h4 = doc.add_heading("4. Post-Remediation Retesting & Verification Audit", level=1)
    h4.style.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    p_retest = doc.add_paragraph()
    p_retest.add_run(
        "Following the engineering of Secure V2 defenses, Team Bug Smasher re-executed all 10 attack scripts against the hardened environment to ensure comprehensive remediation:"
    )

    retest_table = doc.add_table(rows=11, cols=4)
    retest_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    retest_table.autofit = False

    re_headers = ["Vulnerability ID", "Target Vector", "Original V1 Exploit Status", "Secure V2 Verification Result"]
    re_widths = [Inches(1.2), Inches(2.2), Inches(1.3), Inches(1.8)]

    for c_idx, head_text in enumerate(re_headers):
        cell = retest_table.rows[0].cells[c_idx]
        cell.width = re_widths[c_idx]
        set_cell_background(cell, "276749")
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        p = cell.paragraphs[0]
        r = p.add_run(head_text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    retest_rows = [
        ("VULN-001", "Client-Side Route Guarding", "Vulnerable", "BLOCKED (403 Forbidden / RLS Policy Violation)"),
        ("VULN-002", "IDOR on Marks Lookup", "Vulnerable", "BLOCKED (0 Rows Returned; URL parameter removed)"),
        ("VULN-003", "RLS Database Misconfiguration", "Vulnerable", "BLOCKED (Unconstrained query returns only own rows)"),
        ("VULN-004", "Role Privilege Escalation", "Vulnerable", "BLOCKED (Trigger trg_protect_profile_role aborts)"),
        ("VULN-005", "Stored Notice Board XSS", "Vulnerable", "BLOCKED (DOMPurify sanitizes all script/onerror tags)"),
        ("VULN-006", "Mass Assignment Binding", "Vulnerable", "BLOCKED (Whitelist & trigger reject unauthorized keys)"),
        ("VULN-007", "Storage Cross-Tenant Exposure", "Vulnerable", "BLOCKED (Private bucket + 60s signed token required)"),
        ("VULN-008", "PII Sensitive Data Exposure", "Vulnerable", "BLOCKED (Private phone numbers omitted from query)"),
        ("VULN-009", "Attendance Logic Self-Mark", "Vulnerable", "BLOCKED (RLS rejects non-instructor attendance writes)"),
        ("VULN-010", "Client Debug State Leak", "Vulnerable", "BLOCKED (window.__APP_DEBUG__ returns undefined)"),
    ]

    for r_idx, r_data in enumerate(retest_rows):
        row = retest_table.rows[r_idx + 1]
        bg_color = "F0FFF4" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            cell.width = re_widths[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=70, bottom=70, left=70, right=70)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            
            if c_idx == 2:
                r.bold = True
                r.font.color.rgb = RGBColor(0xC5, 0x30, 0x30)
            elif c_idx == 3:
                r.bold = True
                r.font.color.rgb = RGBColor(0x27, 0x67, 0x49)

    doc.add_paragraph()

    # =========================================================================
    # SECTION 5: ADVANCED DEFENSE-IN-DEPTH LESSONS
    # =========================================================================
    h5 = doc.add_heading("5. Advanced Defense-in-Depth Architectural Principles", level=1)
    h5.style.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    p_lessons = doc.add_paragraph()
    p_lessons.add_run(
        "Through this assessment and remediation lifecycle, Team Bug Smasher established 5 core pillars of enterprise web security:\n\n"
        "1. Shift Authorization to the Storage & Database Layer:\n"
        "Client-side UI route guards (such as React Router guards) are purely UX enhancements and must never be treated as access control mechanisms. All authorization checks must be enforced at the database or backend API layer via Row Level Security (RLS).\n\n"
        "2. Granular Tenant Partitioning Over Blanket Policies:\n"
        "Using open clauses like USING (true) in RLS gives a false sense of security while enabling complete table dumps. Every table must evaluate identity claims (e.g. user_id = auth.uid() or get_student_id()).\n\n"
        "3. Database Triggers for Immutable Sensitive Columns:\n"
        "For critical model attributes like user roles and account statuses, RLS policies must be reinforced with BEFORE UPDATE database triggers to eliminate privilege escalation and mass assignment vulnerabilities.\n\n"
        "4. Strict Context-Aware Sanitization:\n"
        "When user-supplied HTML must be rendered, the application must pass input through a hardened sanitizer like DOMPurify with an explicit forbidden tag/attribute list (stripping script, iframe, and inline event handlers).\n\n"
        "5. Ephemeral Signed Access for Multi-Tenant Storage:\n"
        "Cloud storage buckets storing sensitive user data must remain private, using short-lived (60s) cryptographic signed URLs with path-scoped RLS policies."
    )

    p_sign = doc.add_paragraph()
    p_sign.add_run("\nRespectfully Submitted by Team Bug Smasher:\n").bold = True
    p_sign.add_run(
        "• Ranjan Gowda S S — Team Leader & Full-Stack Security Engineer\n"
        "• Manish M — Lead Penetration Tester & Security Researcher\n"
        "• Keerthan Gowda T R — Database & Backend Security Engineer"
    )

    doc.save("c:\\Users\\manis\\OneDrive\\Desktop\\hackathon\\HACKATHON_REPORT_BUG_SMASHER.docx")
    print("SUCCESS: Enriched HACKATHON_REPORT_BUG_SMASHER.docx generated successfully!")

if __name__ == "__main__":
    create_report()
